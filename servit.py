"""
Servit — Mobile Server Dashboard

Manage your Linux server from your phone: terminal (with Korean IME),
file browser/editor, system monitoring, Docker management, log viewer, and notes.

Single dependency: pip install websockets
Optional: openpyxl (Excel), python-docx (DOCX), pyhwp (HWP)

Usage:
    python servit.py                        # System auth, port 8765
    python servit.py --port 9000            # Custom port
    python servit.py --pass mypassword      # Single password mode
    python servit.py --accounts "u1:p1,u2:p2"  # Static accounts

Source: https://github.com/modak000/servit
License: MIT
"""

import asyncio
import base64
import json
import os
import pty
import signal
import struct
import fcntl
import termios
import subprocess
import sys
import stat
import hashlib
import secrets
import select
import pwd
import glob as glob_mod
from pathlib import Path
from http import HTTPStatus

import websockets
from websockets.asyncio.server import serve
from websockets.http11 import Request, Response

STATIC_DIR = Path(__file__).parent / "static"
HOST = "0.0.0.0"
PORT = int(sys.argv[sys.argv.index("--port") + 1]) if "--port" in sys.argv else 8765

# ── Auth ──────────────────────────────────────────────────────────────
# Mode 1 (default): --system-auth  → validate against real Linux passwords via su
# Mode 2: --accounts "user1:pass1,user2:pass2"  → static accounts
# Mode 3: --pass password  → single password, current user
AUTH_MODE = "system"  # system | static | single

ACCOUNTS = {}
FALLBACK_PASSWORD = None
FALLBACK_ROOT = sys.argv[sys.argv.index("--root") + 1] if "--root" in sys.argv else str(Path.home())

if "--accounts" in sys.argv:
    AUTH_MODE = "static"
    raw = sys.argv[sys.argv.index("--accounts") + 1]
    for pair in raw.split(","):
        pair = pair.strip()
        if ":" in pair:
            u, p = pair.split(":", 1)
            ACCOUNTS[u.strip()] = p.strip()
elif "--pass" in sys.argv:
    AUTH_MODE = "single"
    FALLBACK_PASSWORD = sys.argv[sys.argv.index("--pass") + 1]
else:
    AUTH_MODE = "system"

# Active sessions: token -> {username, home}
ACTIVE_TOKENS = {}

# Brute force protection: IP -> {attempts, locked_until}
import time
import re
LOGIN_ATTEMPTS = {}  # ip -> {"count": int, "first": float, "locked_until": float}
MAX_ATTEMPTS = 10        # 10회 실패 허용
ATTEMPT_WINDOW = 300     # 5분 내
LOCKOUT_DURATION = 600   # 10분 잠금

# API rate limiting: IP -> {"count": int, "window_start": float}
API_RATE_LIMITS = {}
API_RATE_MAX = 100       # max requests per minute
API_RATE_WINDOW = 60     # 1 minute window


def check_rate_limit(ip):
    """Returns (allowed, remaining_lockout_seconds)."""
    now = time.time()
    info = LOGIN_ATTEMPTS.get(ip)
    if not info:
        return True, 0

    # 잠금 중인지 확인
    if info.get("locked_until", 0) > now:
        remaining = int(info["locked_until"] - now)
        return False, remaining

    # 윈도우 만료 → 초기화
    if now - info.get("first", 0) > ATTEMPT_WINDOW:
        LOGIN_ATTEMPTS.pop(ip, None)
        return True, 0

    return True, 0


def record_failed_attempt(ip):
    """Record a failed login and lock if threshold reached."""
    now = time.time()
    info = LOGIN_ATTEMPTS.get(ip, {"count": 0, "first": now, "locked_until": 0})

    # 윈도우 만료 → 리셋
    if now - info["first"] > ATTEMPT_WINDOW:
        info = {"count": 0, "first": now, "locked_until": 0}

    info["count"] += 1

    if info["count"] >= MAX_ATTEMPTS:
        info["locked_until"] = now + LOCKOUT_DURATION

    LOGIN_ATTEMPTS[ip] = info


def record_successful_login(ip):
    """Clear failed attempts on success."""
    LOGIN_ATTEMPTS.pop(ip, None)


def check_api_rate_limit(ip):
    """Check API rate limit. Returns (allowed, retry_after_seconds)."""
    now = time.time()
    info = API_RATE_LIMITS.get(ip)
    if not info:
        API_RATE_LIMITS[ip] = {"count": 1, "window_start": now}
        return True, 0
    # Window expired -> reset
    if now - info["window_start"] > API_RATE_WINDOW:
        API_RATE_LIMITS[ip] = {"count": 1, "window_start": now}
        return True, 0
    info["count"] += 1
    if info["count"] > API_RATE_MAX:
        remaining = int(API_RATE_WINDOW - (now - info["window_start"])) + 1
        return False, remaining
    return True, 0


def validate_path_access(requested_path, user_home, allow_log_read=False):
    """Validate that a path is within the user's allowed directories.
    Returns (allowed, resolved_path) tuple."""
    try:
        resolved = os.path.realpath(requested_path)
    except (ValueError, OSError):
        return False, ""

    # Always allow paths under user home
    home_resolved = os.path.realpath(user_home)
    if resolved.startswith(home_resolved + "/") or resolved == home_resolved:
        return True, resolved

    # Allow read-only access to log paths
    if allow_log_read:
        log_dirs = ["/var/log/", "/tmp/servit"]
        for log_dir in log_dirs:
            if resolved.startswith(log_dir):
                return True, resolved
        if resolved == "__journalctl__":
            return True, resolved

    # Block sensitive paths explicitly
    return False, resolved


def validate_docker_id(container_id):
    """Validate Docker container ID/name format to prevent injection."""
    if not container_id:
        return False
    # Docker IDs: hex (short 12 or full 64 chars)
    # Docker names: alphanumeric, underscore, hyphen, dot, slash (for compose)
    if re.match(r'^[a-zA-Z0-9][a-zA-Z0-9_.\-/]{0,127}$', container_id):
        return True
    return False


def make_token(username, password):
    return hashlib.sha256(f"{username}:{password}:{secrets.token_hex(8)}".encode()).hexdigest()[:32]


def verify_system_password(username, password):
    """Verify Linux system password using su + PTY. Works without root."""
    try:
        pw_entry = pwd.getpwnam(username)
    except KeyError:
        return False

    master, slave = pty.openpty()
    pid = os.fork()

    if pid == 0:
        # Child process
        os.close(master)
        os.setsid()
        os.dup2(slave, 0)
        os.dup2(slave, 1)
        os.dup2(slave, 2)
        os.close(slave)
        os.execvp("su", ["su", "-c", "echo __AUTH_OK__", username])
        os._exit(1)

    os.close(slave)

    try:
        # Wait for password prompt
        output = b""
        for _ in range(30):
            r, _, _ = select.select([master], [], [], 0.2)
            if r:
                try:
                    output += os.read(master, 4096)
                except OSError:
                    break
            if b"assword" in output:
                break

        if b"assword" not in output:
            return False

        # Send password
        os.write(master, (password + "\n").encode())

        # Read result
        result = b""
        for _ in range(30):
            r, _, _ = select.select([master], [], [], 0.2)
            if r:
                try:
                    result += os.read(master, 4096)
                except OSError:
                    break
            if b"__AUTH_OK__" in result or b"failure" in result or b"incorrect" in result:
                break

        return b"__AUTH_OK__" in result

    finally:
        try:
            os.close(master)
        except OSError:
            pass
        try:
            os.waitpid(pid, 0)
        except ChildProcessError:
            pass


def authenticate_user(username, password):
    """Authenticate and return (success, home_dir)."""
    if AUTH_MODE == "system":
        if verify_system_password(username, password):
            home = pwd.getpwnam(username).pw_dir
            return True, home
        return False, ""

    elif AUTH_MODE == "static":
        if username in ACCOUNTS and ACCOUNTS[username] == password:
            try:
                home = pwd.getpwnam(username).pw_dir
            except KeyError:
                home = f"/home/{username}"
            return True, home
        return False, ""

    elif AUTH_MODE == "single":
        if password == FALLBACK_PASSWORD:
            return True, FALLBACK_ROOT
        return False, ""

    return False, ""


def precompute_tokens():
    """Precompute tokens for static accounts."""
    if AUTH_MODE == "static":
        for username, password in ACCOUNTS.items():
            token = make_token(username, password)
            try:
                home = pwd.getpwnam(username).pw_dir
            except KeyError:
                home = f"/home/{username}"
            ACTIVE_TOKENS[token] = {"username": username, "home": home}
    elif AUTH_MODE == "single" and FALLBACK_PASSWORD:
        token = make_token("_default_", FALLBACK_PASSWORD)
        ACTIVE_TOKENS[token] = {"username": os.environ.get("USER", "user"), "home": FALLBACK_ROOT}


def get_token_from_cookies(cookies_str):
    """Extract token value from Cookie header."""
    for part in cookies_str.split(";"):
        part = part.strip()
        if part.startswith("token="):
            return part[6:]
    return None


def check_auth(request):
    """Check cookie-based auth token. Returns token string or None."""
    cookies = request.headers.get("Cookie", "")
    token = get_token_from_cookies(cookies)
    if token and token in ACTIVE_TOKENS:
        return token
    return None


def get_session_info(token):
    """Get username and home dir for a token."""
    return ACTIVE_TOKENS.get(token, {"username": "unknown", "home": "/tmp"})


SECURITY_HEADERS = [
    ("X-Content-Type-Options", "nosniff"),
    ("X-Frame-Options", "DENY"),
    ("Content-Security-Policy",
     "default-src 'self' https://cdn.jsdelivr.net; "
     "connect-src 'self' ws: wss:; "
     "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net; "
     "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net"),
]


def make_response(status_code, body, content_type="text/plain", extra_headers=None):
    h = [("Content-Type", content_type), ("Content-Length", str(len(body)))]
    h.extend(SECURITY_HEADERS)
    if extra_headers:
        h.extend(extra_headers)
    return Response(status_code, "OK", websockets.datastructures.Headers(h), body)


def make_json(data, extra_headers=None):
    # Handle surrogate characters from HWP/binary parsing
    text = json.dumps(data, ensure_ascii=False)
    body = text.encode("utf-8", errors="replace")
    return make_response(200, body, "application/json; charset=utf-8", extra_headers)


# ── HTTP Handler ──────────────────────────────────────────────────────
async def process_request(connection, request):
    if request.headers.get("Upgrade", "").lower() == "websocket":
        # WebSocket auth check via query param
        qs = request.path.split("?")[-1] if "?" in request.path else ""
        import urllib.parse
        params = urllib.parse.parse_qs(qs)
        token = params.get("token", [""])[0]
        if token and token in ACTIVE_TOKENS:
            # Stash session info on the connection for later use
            connection.session_info = ACTIVE_TOKENS[token]
            return None
        # Also check cookie
        cookie_token = check_auth(request)
        if cookie_token:
            connection.session_info = ACTIVE_TOKENS[cookie_token]
            return None
        return make_response(403, b"Forbidden")

    path = request.path.split("?")[0]

    # Login page (no auth needed)
    if path == "/login":
        return make_response(200, (STATIC_DIR / "login.html").read_bytes(), "text/html; charset=utf-8")

    # Login — credentials via Authorization: Basic header (never in URL)
    if path == "/api/login":
        username = ""
        pw = ""

        # Parse credentials from Authorization: Basic header
        auth_header = request.headers.get("Authorization", "")
        if auth_header.startswith("Basic "):
            try:
                decoded = base64.b64decode(auth_header[6:]).decode("utf-8")
                if ":" in decoded:
                    username, pw = decoded.split(":", 1)
                    username = username.strip()
            except Exception:
                pass

        # Get client IP (Cloudflare passes real IP in headers)
        client_ip = (
            request.headers.get("CF-Connecting-IP")
            or request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
            or "unknown"
        )

        # Rate limit check
        allowed, lockout_remaining = check_rate_limit(client_ip)
        if not allowed:
            return make_json({
                "ok": False,
                "error": f"Too many attempts. Try again in {lockout_remaining // 60} min {lockout_remaining % 60} sec"
            })

        if not username and AUTH_MODE == "single":
            username = "_default_"

        if not username or not pw:
            return make_json({"ok": False, "error": "Username and password required"})

        success, home = authenticate_user(username, pw)
        if success:
            record_successful_login(client_ip)
            token = make_token(username, pw)
            ACTIVE_TOKENS[token] = {"username": username, "home": home}
            return make_json(
                {"ok": True, "token": token, "username": username, "home": home},
                [("Set-Cookie", f"token={token}; Path=/; HttpOnly; SameSite=Strict")]
            )

        record_failed_attempt(client_ip)
        remaining = MAX_ATTEMPTS - LOGIN_ATTEMPTS.get(client_ip, {}).get("count", 0)
        error_msg = "Invalid username or password"
        if remaining <= 3:
            error_msg += f" ({remaining} attempts remaining)"
        return make_json({"ok": False, "error": error_msg})

    # All other routes need auth
    if not check_auth(request):
        return make_response(302, b"", "text/plain", [("Location", "/login")])

    if path == "/" or path == "":
        path = "/index.html"

    # API — need to pass session info + rate limiting
    if path.startswith("/api/"):
        client_ip = (
            request.headers.get("CF-Connecting-IP")
            or request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
            or "unknown"
        )
        api_allowed, retry_after = check_api_rate_limit(client_ip)
        if not api_allowed:
            return make_json({"error": f"Rate limit exceeded. Retry after {retry_after}s"})

        token = check_auth(request)
        info = get_session_info(token) if token else {"username": "unknown", "home": "/tmp"}
        return handle_api(path, request, info)

    # Static files
    file_path = STATIC_DIR / path.lstrip("/")
    if not file_path.exists() or not file_path.is_file():
        return make_response(404, b"404")

    ct_map = {
        ".html": "text/html; charset=utf-8",
        ".js": "application/javascript; charset=utf-8",
        ".css": "text/css; charset=utf-8",
        ".svg": "image/svg+xml",
        ".json": "application/json; charset=utf-8",
    }
    return make_response(200, file_path.read_bytes(), ct_map.get(file_path.suffix, "application/octet-stream"))


def collect_stats():
    """Collect server stats from /proc and system tools."""
    stats = {}

    # CPU — load average (instant, no delay needed)
    try:
        load = os.getloadavg()
        stats["load"] = list(load)
        # Count logical CPUs
        cpu_count = os.cpu_count() or 1
        stats["cpu"] = {
            "percent": round(load[0] / cpu_count * 100, 1),
            "cores": cpu_count,
        }
    except OSError:
        stats["cpu"] = {"percent": 0, "cores": 1}
        stats["load"] = [0, 0, 0]

    # Memory — /proc/meminfo
    try:
        meminfo = {}
        with open("/proc/meminfo", "r") as f:
            for line in f:
                parts = line.split()
                if len(parts) >= 2:
                    key = parts[0].rstrip(":")
                    val = int(parts[1]) * 1024  # kB to bytes
                    meminfo[key] = val
        total = meminfo.get("MemTotal", 0)
        available = meminfo.get("MemAvailable", 0)
        used = total - available
        pct = round(used / total * 100, 1) if total else 0
        stats["memory"] = {"total": total, "used": used, "percent": pct}
    except Exception:
        stats["memory"] = {"total": 0, "used": 0, "percent": 0}

    # Disk — os.statvfs
    try:
        sv = os.statvfs("/")
        total = sv.f_frsize * sv.f_blocks
        free = sv.f_frsize * sv.f_bavail
        used = total - free
        pct = round(used / total * 100, 1) if total else 0
        stats["disk"] = {"total": total, "used": used, "percent": pct, "path": "/"}
    except Exception:
        stats["disk"] = {"total": 0, "used": 0, "percent": 0, "path": "/"}

    # GPU — nvidia-smi
    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=name,utilization.gpu,memory.used,memory.total,temperature.gpu",
             "--format=csv,nounits,noheader"],
            capture_output=True, text=True, timeout=5
        )
        gpus = []
        for line in result.stdout.strip().split("\n"):
            if not line.strip():
                continue
            parts = [p.strip() for p in line.split(",")]
            if len(parts) >= 5:
                gpus.append({
                    "name": parts[0],
                    "util": int(parts[1]) if parts[1].isdigit() else 0,
                    "memory_used": int(parts[2]) if parts[2].isdigit() else 0,
                    "memory_total": int(parts[3]) if parts[3].isdigit() else 0,
                    "temp": int(parts[4]) if parts[4].isdigit() else 0,
                })
        stats["gpu"] = gpus
    except Exception:
        stats["gpu"] = []

    # Network — /proc/net/dev
    try:
        with open("/proc/net/dev", "r") as f:
            lines = f.readlines()
        total_recv = 0
        total_sent = 0
        for line in lines[2:]:
            parts = line.split()
            if len(parts) >= 10:
                iface = parts[0].rstrip(":")
                if iface == "lo":
                    continue
                total_recv += int(parts[1])
                total_sent += int(parts[9])
        stats["network"] = {"bytes_sent": total_sent, "bytes_recv": total_recv}
    except Exception:
        stats["network"] = {"bytes_sent": 0, "bytes_recv": 0}

    # Uptime — /proc/uptime
    try:
        with open("/proc/uptime", "r") as f:
            uptime_secs = float(f.read().split()[0])
        days = int(uptime_secs // 86400)
        hours = int((uptime_secs % 86400) // 3600)
        minutes = int((uptime_secs % 3600) // 60)
        if days > 0:
            stats["uptime"] = f"{days} days, {hours}:{minutes:02d}"
        else:
            stats["uptime"] = f"{hours}:{minutes:02d}"
    except Exception:
        stats["uptime"] = "unknown"

    # Processes — count /proc/[0-9]*/
    try:
        stats["processes"] = len(glob_mod.glob("/proc/[0-9]*/"))
    except Exception:
        stats["processes"] = 0

    return stats


def handle_api(path, request, session_info):
    import urllib.parse
    qs = request.path.split("?")[-1] if "?" in request.path else ""
    params = urllib.parse.parse_qs(qs)
    user_root = session_info["home"]

    if path.startswith("/api/tree"):
        dir_path = params.get("path", [user_root])[0]
        if not dir_path or dir_path == "":
            dir_path = user_root
        if not os.path.isabs(dir_path):
            dir_path = os.path.join(user_root, dir_path)

        allowed, resolved = validate_path_access(dir_path, user_root)
        if not allowed:
            return make_json({"error": "Access denied: path outside home directory", "path": dir_path, "entries": []})
        dir_path = resolved

        show_hidden = params.get("hidden", ["0"])[0] == "1"
        entries = []
        try:
            for name in sorted(os.listdir(dir_path)):
                if not show_hidden and name.startswith("."):
                    continue
                full = os.path.join(dir_path, name)
                try:
                    s = os.stat(full, follow_symlinks=False)
                    entries.append({
                        "name": name,
                        "path": full,
                        "is_dir": stat.S_ISDIR(s.st_mode),
                        "is_link": stat.S_ISLNK(os.lstat(full).st_mode) if os.path.islink(full) else False,
                        "size": s.st_size,
                        "ext": Path(name).suffix.lower(),
                    })
                except (PermissionError, OSError):
                    pass
        except (PermissionError, OSError) as e:
            return make_json({"error": str(e), "path": dir_path, "entries": []})

        entries.sort(key=lambda x: (not x["is_dir"], x["name"].lower()))
        parent = str(Path(dir_path).parent)
        return make_json({"path": dir_path, "parent": parent, "entries": entries, "home": user_root})

    elif path.startswith("/api/file"):
        file_path = params.get("path", [""])[0]
        if not file_path or not os.path.isfile(file_path):
            return make_json({"error": "Not found"})

        allowed, resolved = validate_path_access(file_path, user_root)
        if not allowed:
            return make_json({"error": "Access denied: path outside home directory"})
        file_path = resolved

        size = os.path.getsize(file_path)
        ext = Path(file_path).suffix.lower()

        # Office documents — parse server-side
        if ext in {".xlsx", ".xls"}:
            return parse_excel(file_path, size)
        if ext == ".hwp":
            return parse_hwp(file_path, size)
        if ext == ".docx":
            return parse_docx(file_path, size)

        # Binary check
        binary_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".ico", ".pdf",
                       ".zip", ".tar", ".gz", ".bz2", ".xz", ".7z",
                       ".so", ".o", ".pyc", ".class", ".jar", ".ttf", ".woff2",
                       ".mp3", ".mp4", ".wav", ".avi", ".mkv", ".pptx"}
        if ext in binary_exts:
            return make_json({
                "path": file_path,
                "content": f"[Binary file -- {ext} / {size:,} bytes]",
                "language": "text", "size": size, "binary": True
            })

        if size > 500_000:
            return make_json({
                "path": file_path,
                "content": f"[File size {size:,} bytes -- exceeds 500KB limit]",
                "language": "text", "size": size, "truncated": True
            })

        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
        except Exception as e:
            content = f"[Read failed: {e}]"

        lang_map = {
            ".py": "python", ".js": "javascript", ".ts": "typescript",
            ".kt": "kotlin", ".java": "java", ".sh": "bash",
            ".yml": "yaml", ".yaml": "yaml", ".json": "json",
            ".md": "markdown", ".html": "html", ".css": "css",
            ".xml": "xml", ".gradle": "groovy", ".txt": "text",
            ".csv": "text", ".log": "text", ".properties": "properties",
            ".toml": "toml", ".cfg": "text", ".conf": "text",
            ".rs": "rust", ".go": "go", ".c": "c", ".h": "c",
            ".cpp": "cpp", ".hpp": "cpp", ".sql": "sql",
        }
        return make_json({
            "path": file_path, "content": content,
            "language": lang_map.get(ext, "text"), "size": size,
        })

    elif path.startswith("/api/search"):
        query = params.get("q", [""])[0]
        search_in = params.get("path", [user_root])[0]
        if not query:
            return make_json({"results": []})

        allowed, resolved = validate_path_access(search_in, user_root)
        if not allowed:
            return make_json({"results": [], "error": "Access denied: path outside home directory"})
        search_in = resolved

        results = []
        try:
            proc = subprocess.run(
                ["grep", "-rn", "--include=*.py", "--include=*.kt",
                 "--include=*.java", "--include=*.js", "--include=*.md",
                 "--include=*.yml", "--include=*.yaml", "--include=*.sh",
                 "--include=*.txt", "--include=*.json", "--include=*.xml",
                 "--include=*.html", "--include=*.css", "--include=*.gradle",
                 "-m", "3", "-l", query, search_in],
                capture_output=True, text=True, timeout=5
            )
            for line in proc.stdout.strip().split("\n")[:30]:
                if line:
                    results.append(line)
        except Exception:
            pass
        return make_json({"query": query, "results": results})

    elif path.startswith("/api/stats"):
        return make_json(collect_stats())

    elif path.startswith("/api/save"):
        file_path = params.get("path", [""])[0]
        b64_content = params.get("b64", [""])[0]

        if not file_path:
            return make_json({"ok": False, "error": "No path specified"})

        if not b64_content:
            return make_json({"ok": False, "error": "No content provided"})

        if not os.path.isfile(file_path):
            return make_json({"ok": False, "error": "File not found"})

        allowed, resolved = validate_path_access(file_path, user_root)
        if not allowed:
            return make_json({"ok": False, "error": "Access denied: path outside home directory"})
        file_path = resolved

        if not os.access(file_path, os.W_OK):
            return make_json({"ok": False, "error": "File is not writable"})

        try:
            content = base64.b64decode(b64_content).decode("utf-8")
            # Create backup before overwrite (single .bak copy)
            backup_path = file_path + ".servit-bak"
            try:
                import shutil
                shutil.copy2(file_path, backup_path)
            except Exception:
                pass  # Best effort backup
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            return make_json({"ok": True})
        except Exception as e:
            return make_json({"ok": False, "error": str(e)})

    elif path.startswith("/api/docker"):
        action = params.get("action", ["list"])[0]
        try:
            if action == "list":
                result = subprocess.run(
                    ["docker", "ps", "-a", "--format", '{{json .}}'],
                    capture_output=True, text=True, timeout=10
                )
                if result.returncode != 0 and "not found" in (result.stderr or "").lower():
                    return make_json({"error": "Docker not installed"})
                containers = []
                for line in result.stdout.strip().split("\n"):
                    if not line.strip():
                        continue
                    try:
                        c = json.loads(line)
                        containers.append({
                            "name": c.get("Names", ""),
                            "status": c.get("Status", ""),
                            "image": c.get("Image", ""),
                            "ports": c.get("Ports", ""),
                            "id": c.get("ID", ""),
                            "state": c.get("State", ""),
                        })
                    except json.JSONDecodeError:
                        pass
                return make_json({"containers": containers})
            elif action == "logs":
                cid = params.get("id", [""])[0]
                lines = params.get("lines", ["100"])[0]
                if not cid:
                    return make_json({"error": "No container ID"})
                if not validate_docker_id(cid):
                    return make_json({"error": "Invalid container ID format"})
                # Validate lines is a number
                if not lines.isdigit() or int(lines) > 10000:
                    lines = "100"
                result = subprocess.run(
                    ["docker", "logs", "--tail", lines, cid],
                    capture_output=True, text=True, timeout=10
                )
                return make_json({"logs": result.stdout + result.stderr})
            elif action in ("start", "stop", "restart"):
                cid = params.get("id", [""])[0]
                if not cid:
                    return make_json({"error": "No container ID"})
                if not validate_docker_id(cid):
                    return make_json({"error": "Invalid container ID format"})
                result = subprocess.run(
                    ["docker", action, cid],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0:
                    return make_json({"ok": True, "message": f"Container {action}ed"})
                return make_json({"ok": False, "error": result.stderr.strip()})
            else:
                return make_json({"error": "Unknown action"})
        except FileNotFoundError:
            return make_json({"error": "Docker not installed"})
        except subprocess.TimeoutExpired:
            return make_json({"error": "Docker command timed out"})
        except Exception as e:
            return make_json({"error": str(e)})

    elif path.startswith("/api/processes"):
        try:
            result = subprocess.run(
                ["ps", "aux", "--sort=-%cpu"],
                capture_output=True, text=True, timeout=5
            )
            lines = result.stdout.strip().split("\n")
            processes = []
            my_pid = os.getpid()
            for line in lines[1:31]:  # skip header, top 30
                parts = line.split(None, 10)
                if len(parts) >= 11:
                    processes.append({
                        "pid": int(parts[1]),
                        "user": parts[0],
                        "cpu": float(parts[2]),
                        "mem": float(parts[3]),
                        "command": parts[10][:200],
                    })
            return make_json({"processes": processes, "servit_pid": my_pid})
        except Exception as e:
            return make_json({"error": str(e)})

    elif path.startswith("/api/kill"):
        try:
            pid = int(params.get("pid", ["0"])[0])
            sig = int(params.get("signal", ["15"])[0])
            if pid <= 1:
                return make_json({"ok": False, "error": "Cannot kill PID 0 or 1"})
            if pid == os.getpid():
                return make_json({"ok": False, "error": "Cannot kill Servit process"})
            if sig not in (9, 15):
                return make_json({"ok": False, "error": "Only SIGTERM(15) and SIGKILL(9) allowed"})

            # Restrict to processes owned by the authenticated user
            try:
                with open(f"/proc/{pid}/status", "r") as f:
                    proc_status = f.read()
                proc_uid = None
                for line in proc_status.split("\n"):
                    if line.startswith("Uid:"):
                        # Format: Uid:\treal\teffective\tsaved\tfs
                        proc_uid = int(line.split()[1])
                        break
                if proc_uid is None:
                    return make_json({"ok": False, "error": "Cannot determine process owner"})
                current_uid = os.getuid()
                if proc_uid != current_uid:
                    return make_json({"ok": False, "error": "Cannot kill processes owned by other users"})
            except FileNotFoundError:
                return make_json({"ok": False, "error": "Process not found"})
            except (PermissionError, OSError):
                return make_json({"ok": False, "error": "Cannot verify process owner"})

            os.kill(pid, sig)
            return make_json({"ok": True})
        except ProcessLookupError:
            return make_json({"ok": False, "error": "Process not found"})
        except PermissionError:
            return make_json({"ok": False, "error": "Permission denied"})
        except Exception as e:
            return make_json({"ok": False, "error": str(e)})

    elif path.startswith("/api/logs"):
        action = params.get("action", ["list"])[0]
        if action == "list":
            log_files = []
            search_paths = ["/var/log/*.log", "/var/log/syslog", "/var/log/auth.log",
                            os.path.expanduser("~/.local/share/**/*.log"),
                            "/tmp/servit*.log"]
            for pattern in search_paths:
                for f in glob_mod.glob(pattern, recursive=True):
                    try:
                        st = os.stat(f)
                        log_files.append({
                            "name": os.path.basename(f),
                            "path": f,
                            "size": st.st_size,
                            "modified": st.st_mtime,
                        })
                    except (PermissionError, OSError):
                        pass
            # Also try journalctl units
            try:
                result = subprocess.run(
                    ["journalctl", "--list-boots", "--no-pager"],
                    capture_output=True, text=True, timeout=3
                )
                if result.returncode == 0:
                    log_files.append({
                        "name": "journalctl (system)",
                        "path": "__journalctl__",
                        "size": 0,
                        "modified": time.time(),
                    })
            except Exception:
                pass
            log_files.sort(key=lambda x: x["modified"], reverse=True)
            return make_json({"files": log_files})
        elif action == "read":
            log_path = params.get("path", [""])[0]
            lines_count = params.get("lines", ["200"])[0]
            if not log_path:
                return make_json({"error": "No path specified"})
            # Validate log path access
            allowed, _ = validate_path_access(log_path, user_root, allow_log_read=True)
            if not allowed:
                return make_json({"error": "Access denied: path not in allowed log directories"})
            try:
                if log_path == "__journalctl__":
                    result = subprocess.run(
                        ["journalctl", "-n", lines_count, "--no-pager"],
                        capture_output=True, text=True, timeout=5
                    )
                    return make_json({"content": result.stdout})
                result = subprocess.run(
                    ["tail", "-n", lines_count, log_path],
                    capture_output=True, text=True, timeout=5
                )
                if result.returncode != 0:
                    return make_json({"error": result.stderr.strip() or "Failed to read log"})
                return make_json({"content": result.stdout})
            except PermissionError:
                return make_json({"error": "Permission denied"})
            except Exception as e:
                return make_json({"error": str(e)})
        elif action == "search":
            log_path = params.get("path", [""])[0]
            query = params.get("query", [""])[0]
            if not log_path or not query:
                return make_json({"error": "Path and query required"})
            # Validate log path access
            allowed, _ = validate_path_access(log_path, user_root, allow_log_read=True)
            if not allowed:
                return make_json({"error": "Access denied: path not in allowed log directories"})
            try:
                if log_path == "__journalctl__":
                    result = subprocess.run(
                        ["journalctl", "--no-pager", "-n", "1000", "--grep", query],
                        capture_output=True, text=True, timeout=5
                    )
                    lines = result.stdout.strip().split("\n")[-50:]
                    return make_json({"matches": "\n".join(lines)})
                result = subprocess.run(
                    ["grep", "-n", query, log_path],
                    capture_output=True, text=True, timeout=5
                )
                lines = result.stdout.strip().split("\n")[-50:]
                return make_json({"matches": "\n".join(lines)})
            except Exception as e:
                return make_json({"error": str(e)})
        return make_json({"error": "Unknown action"})

    elif path.startswith("/api/notes"):
        import urllib.parse as _up
        notes_dir = os.path.join(session_info["home"], ".servit", "notes")
        os.makedirs(notes_dir, exist_ok=True)
        action = params.get("action", ["list"])[0]
        if action == "list":
            notes = []
            for f in sorted(os.listdir(notes_dir)):
                if f.endswith(".md"):
                    fp = os.path.join(notes_dir, f)
                    try:
                        st = os.stat(fp)
                        notes.append({
                            "name": f[:-3],  # strip .md
                            "size": st.st_size,
                            "modified": st.st_mtime,
                        })
                    except OSError:
                        pass
            return make_json({"notes": notes})
        elif action == "read":
            name = params.get("name", [""])[0]
            if not name:
                return make_json({"error": "No name specified"})
            fp = os.path.join(notes_dir, name + ".md")
            if not os.path.isfile(fp):
                return make_json({"error": "Note not found"})
            try:
                with open(fp, "r", encoding="utf-8") as f:
                    content = f.read()
                return make_json({"name": name, "content": content})
            except Exception as e:
                return make_json({"error": str(e)})
        elif action == "save":
            name = params.get("name", [""])[0]
            b64_content = params.get("b64", [""])[0]
            if not name:
                return make_json({"ok": False, "error": "No name specified"})
            # Sanitize name
            safe_name = "".join(c for c in name if c.isalnum() or c in "-_ ").strip()
            if not safe_name:
                return make_json({"ok": False, "error": "Invalid name"})
            fp = os.path.join(notes_dir, safe_name + ".md")
            try:
                content = base64.b64decode(b64_content).decode("utf-8") if b64_content else ""
                with open(fp, "w", encoding="utf-8") as f:
                    f.write(content)
                return make_json({"ok": True})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        elif action == "delete":
            name = params.get("name", [""])[0]
            if not name:
                return make_json({"ok": False, "error": "No name specified"})
            fp = os.path.join(notes_dir, name + ".md")
            if not os.path.isfile(fp):
                return make_json({"ok": False, "error": "Note not found"})
            try:
                os.remove(fp)
                return make_json({"ok": True})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        elif action == "create":
            name = params.get("name", [""])[0]
            if not name:
                return make_json({"ok": False, "error": "No name specified"})
            safe_name = "".join(c for c in name if c.isalnum() or c in "-_ ").strip()
            if not safe_name:
                return make_json({"ok": False, "error": "Invalid name"})
            fp = os.path.join(notes_dir, safe_name + ".md")
            if os.path.exists(fp):
                return make_json({"ok": False, "error": "Note already exists"})
            try:
                with open(fp, "w", encoding="utf-8") as f:
                    f.write("")
                return make_json({"ok": True, "name": safe_name})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        return make_json({"error": "Unknown action"})

    elif path.startswith("/api/session"):
        return make_json({
            "username": session_info["username"],
            "home": session_info["home"],
        })

    elif path.startswith("/api/cron"):
        action = params.get("action", ["list"])[0]
        if action == "list":
            try:
                result = subprocess.run(
                    ["crontab", "-l"],
                    capture_output=True, text=True, timeout=5
                )
                lines = result.stdout.strip().split("\n") if result.stdout.strip() else []
                crons = []
                for i, line in enumerate(lines):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    is_comment = stripped.startswith("#")
                    crons.append({"index": i, "line": line, "comment": is_comment})
                return make_json({"crons": crons, "raw": result.stdout})
            except Exception as e:
                return make_json({"error": str(e)})
        elif action == "save":
            raw = params.get("b64", [""])[0]
            if not raw:
                return make_json({"ok": False, "error": "No content"})
            try:
                content = base64.b64decode(raw).decode("utf-8")
                proc = subprocess.run(
                    ["crontab", "-"],
                    input=content, capture_output=True, text=True, timeout=5
                )
                if proc.returncode == 0:
                    return make_json({"ok": True})
                return make_json({"ok": False, "error": proc.stderr.strip()})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        elif action == "delete":
            idx = int(params.get("index", ["-1"])[0])
            try:
                result = subprocess.run(["crontab", "-l"], capture_output=True, text=True, timeout=5)
                lines = result.stdout.split("\n")
                if 0 <= idx < len(lines):
                    lines.pop(idx)
                    new_content = "\n".join(lines)
                    proc = subprocess.run(["crontab", "-"], input=new_content, capture_output=True, text=True, timeout=5)
                    if proc.returncode == 0:
                        return make_json({"ok": True})
                    return make_json({"ok": False, "error": proc.stderr.strip()})
                return make_json({"ok": False, "error": "Invalid index"})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        elif action == "add":
            entry = params.get("entry", [""])[0]
            if not entry:
                return make_json({"ok": False, "error": "No entry"})
            try:
                result = subprocess.run(["crontab", "-l"], capture_output=True, text=True, timeout=5)
                current = result.stdout if result.returncode == 0 else ""
                if current and not current.endswith("\n"):
                    current += "\n"
                current += entry + "\n"
                proc = subprocess.run(["crontab", "-"], input=current, capture_output=True, text=True, timeout=5)
                if proc.returncode == 0:
                    return make_json({"ok": True})
                return make_json({"ok": False, "error": proc.stderr.strip()})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        return make_json({"error": "Unknown action"})

    elif path.startswith("/api/diskusage"):
        target = params.get("path", [user_root])[0]
        allowed, resolved = validate_path_access(target, user_root)
        if not allowed:
            return make_json({"error": "Access denied"})
        depth = params.get("depth", ["1"])[0]
        if not depth.isdigit() or int(depth) > 3:
            depth = "1"
        try:
            result = subprocess.run(
                ["du", "-h", "--max-depth=" + depth, "-t", "1M", resolved],
                capture_output=True, text=True, timeout=15
            )
            entries = []
            for line in result.stdout.strip().split("\n"):
                if not line.strip():
                    continue
                parts = line.split("\t", 1)
                if len(parts) == 2:
                    entries.append({"size": parts[0].strip(), "path": parts[1].strip()})
            # Sort by size descending (parse human-readable)
            def parse_size(s):
                s = s.strip()
                try:
                    if s.endswith("G"):
                        return float(s[:-1]) * 1073741824
                    if s.endswith("M"):
                        return float(s[:-1]) * 1048576
                    if s.endswith("K"):
                        return float(s[:-1]) * 1024
                    if s.endswith("T"):
                        return float(s[:-1]) * 1099511627776
                    return float(s)
                except ValueError:
                    return 0
            entries.sort(key=lambda x: parse_size(x["size"]), reverse=True)
            return make_json({"entries": entries[:50], "path": resolved})
        except subprocess.TimeoutExpired:
            return make_json({"error": "Scan timed out (directory too large)"})
        except Exception as e:
            return make_json({"error": str(e)})

    elif path.startswith("/api/netstat"):
        try:
            result = subprocess.run(
                ["ss", "-tunap"],
                capture_output=True, text=True, timeout=5
            )
            connections = []
            for line in result.stdout.strip().split("\n")[1:]:
                parts = line.split()
                if len(parts) >= 5:
                    conn = {
                        "proto": parts[0],
                        "state": parts[1],
                        "local": parts[4] if len(parts) > 4 else "",
                        "peer": parts[5] if len(parts) > 5 else "",
                        "process": parts[6] if len(parts) > 6 else "",
                    }
                    connections.append(conn)
            # Also get listening ports
            result2 = subprocess.run(
                ["ss", "-tlnp"],
                capture_output=True, text=True, timeout=5
            )
            listeners = []
            for line in result2.stdout.strip().split("\n")[1:]:
                parts = line.split()
                if len(parts) >= 4:
                    listeners.append({
                        "proto": parts[0],
                        "local": parts[3] if len(parts) > 3 else "",
                        "process": parts[5] if len(parts) > 5 else parts[-1],
                    })
            return make_json({"connections": connections[:100], "listeners": listeners[:50]})
        except FileNotFoundError:
            return make_json({"error": "ss command not found"})
        except Exception as e:
            return make_json({"error": str(e)})

    elif path.startswith("/api/services"):
        action = params.get("action", ["list"])[0]
        if action == "list":
            try:
                result = subprocess.run(
                    ["systemctl", "list-units", "--type=service", "--all", "--no-pager", "--plain", "--no-legend"],
                    capture_output=True, text=True, timeout=10
                )
                services = []
                for line in result.stdout.strip().split("\n"):
                    if not line.strip():
                        continue
                    parts = line.split(None, 4)
                    if len(parts) >= 4:
                        services.append({
                            "name": parts[0].replace(".service", ""),
                            "load": parts[1],
                            "active": parts[2],
                            "sub": parts[3],
                            "description": parts[4] if len(parts) > 4 else "",
                        })
                return make_json({"services": services})
            except FileNotFoundError:
                return make_json({"error": "systemctl not found"})
            except Exception as e:
                return make_json({"error": str(e)})
        elif action in ("start", "stop", "restart", "status"):
            svc = params.get("name", [""])[0]
            if not svc or not re.match(r'^[a-zA-Z0-9@._-]+$', svc):
                return make_json({"error": "Invalid service name"})
            try:
                result = subprocess.run(
                    ["systemctl", action, svc],
                    capture_output=True, text=True, timeout=15
                )
                if action == "status":
                    return make_json({"output": result.stdout + result.stderr})
                if result.returncode == 0:
                    return make_json({"ok": True, "message": f"Service {action}: {svc}"})
                return make_json({"ok": False, "error": result.stderr.strip() or "Command failed"})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        return make_json({"error": "Unknown action"})

    elif path.startswith("/api/upload"):
        # File upload via base64 in query params
        target_dir = params.get("path", [user_root])[0]
        filename = params.get("name", [""])[0]
        b64_data = params.get("b64", [""])[0]
        if not filename or not b64_data:
            return make_json({"ok": False, "error": "Missing filename or data"})
        # Sanitize filename
        safe_name = os.path.basename(filename)
        if not safe_name or safe_name.startswith("."):
            return make_json({"ok": False, "error": "Invalid filename"})
        allowed, resolved_dir = validate_path_access(target_dir, user_root)
        if not allowed:
            return make_json({"ok": False, "error": "Access denied"})
        full_path = os.path.join(resolved_dir, safe_name)
        if os.path.exists(full_path):
            return make_json({"ok": False, "error": "File already exists"})
        try:
            data = base64.b64decode(b64_data)
            with open(full_path, "wb") as f:
                f.write(data)
            return make_json({"ok": True, "path": full_path, "size": len(data)})
        except Exception as e:
            return make_json({"ok": False, "error": str(e)})

    elif path.startswith("/api/download"):
        file_path = params.get("path", [""])[0]
        if not file_path or not os.path.isfile(file_path):
            return make_json({"error": "File not found"})
        allowed, resolved = validate_path_access(file_path, user_root)
        if not allowed:
            return make_json({"error": "Access denied"})
        size = os.path.getsize(resolved)
        if size > 50_000_000:  # 50MB limit
            return make_json({"error": "File too large (max 50MB)"})
        try:
            with open(resolved, "rb") as f:
                data = f.read()
            fname = os.path.basename(resolved)
            return make_response(200, data, "application/octet-stream",
                                 [("Content-Disposition", f'attachment; filename="{fname}"')])
        except Exception as e:
            return make_json({"error": str(e)})

    elif path.startswith("/api/ssh-servers"):
        servers_file = os.path.join(session_info["home"], ".servit", "servers.json")
        os.makedirs(os.path.dirname(servers_file), exist_ok=True)
        action = params.get("action", ["list"])[0]

        def load_servers():
            try:
                if os.path.isfile(servers_file):
                    with open(servers_file, "r") as f:
                        return json.load(f)
                return []
            except Exception:
                return []

        def save_servers(servers):
            with open(servers_file, "w") as f:
                json.dump(servers, f, indent=2)

        if action == "list":
            return make_json({"servers": load_servers()})

        elif action == "save":
            name = params.get("name", [""])[0]
            host = params.get("host", [""])[0]
            port = params.get("port", ["22"])[0]
            username = params.get("username", [""])[0]
            auth = params.get("auth", ["key"])[0]
            key_path = params.get("key_path", [""])[0]
            jump_host = params.get("jump_host", [""])[0]

            if not name or not host or not username:
                return make_json({"ok": False, "error": "Name, host, and username are required"})
            if not port.isdigit() or int(port) < 1 or int(port) > 65535:
                return make_json({"ok": False, "error": "Invalid port number"})
            # Validate no shell metacharacters in critical fields
            for field_name, field_val in [("host", host), ("username", username), ("jump_host", jump_host)]:
                if field_val and re.search(r'[;&|`$(){}\\"\'\n\r]', field_val):
                    return make_json({"ok": False, "error": f"Invalid characters in {field_name}"})

            servers = load_servers()
            if len(servers) >= 50:
                return make_json({"ok": False, "error": "Maximum 50 servers"})

            new_server = {
                "name": name.strip(),
                "host": host.strip(),
                "port": int(port),
                "username": username.strip(),
                "auth": auth if auth in ("key", "password") else "key",
                "key_path": key_path.strip(),
                "jump_host": jump_host.strip(),
            }

            # Update existing or append
            found = False
            for i, s in enumerate(servers):
                if s.get("name") == name.strip():
                    servers[i] = new_server
                    found = True
                    break
            if not found:
                servers.append(new_server)

            try:
                save_servers(servers)
                return make_json({"ok": True})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})

        elif action == "delete":
            name = params.get("name", [""])[0]
            if not name:
                return make_json({"ok": False, "error": "No name specified"})
            servers = load_servers()
            servers = [s for s in servers if s.get("name") != name]
            try:
                save_servers(servers)
                return make_json({"ok": True})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})

        elif action == "connect":
            name = params.get("name", [""])[0]
            if not name:
                return make_json({"error": "No name specified"})
            servers = load_servers()
            server = next((s for s in servers if s.get("name") == name), None)
            if not server:
                return make_json({"error": "Server not found"})

            # Build SSH command
            parts = ["ssh"]
            if server.get("jump_host"):
                parts.extend(["-J", server["jump_host"]])
            parts.extend(["-o", "StrictHostKeyChecking=accept-new"])
            if server.get("key_path"):
                parts.extend(["-i", server["key_path"]])
            if server.get("port") and server["port"] != 22:
                parts.extend(["-p", str(server["port"])])
            parts.append(f"{server['username']}@{server['host']}")

            return make_json({"command": " ".join(parts), "server": server})

        return make_json({"error": "Unknown action"})

    elif path.startswith("/api/bookmarks"):
        bm_file = os.path.join(session_info["home"], ".servit", "bookmarks.json")
        os.makedirs(os.path.dirname(bm_file), exist_ok=True)
        action = params.get("action", ["list"])[0]
        if action == "list":
            try:
                if os.path.isfile(bm_file):
                    with open(bm_file, "r") as f:
                        bookmarks = json.load(f)
                else:
                    bookmarks = []
                return make_json({"bookmarks": bookmarks})
            except Exception:
                return make_json({"bookmarks": []})
        elif action == "add":
            bm_path = params.get("path", [""])[0]
            bm_name = params.get("name", [""])[0]
            if not bm_path:
                return make_json({"ok": False, "error": "No path"})
            try:
                bookmarks = []
                if os.path.isfile(bm_file):
                    with open(bm_file, "r") as f:
                        bookmarks = json.load(f)
                if len(bookmarks) >= 50:
                    return make_json({"ok": False, "error": "Max 50 bookmarks"})
                bookmarks.append({"path": bm_path, "name": bm_name or os.path.basename(bm_path)})
                with open(bm_file, "w") as f:
                    json.dump(bookmarks, f)
                return make_json({"ok": True})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        elif action == "remove":
            idx = int(params.get("index", ["-1"])[0])
            try:
                bookmarks = []
                if os.path.isfile(bm_file):
                    with open(bm_file, "r") as f:
                        bookmarks = json.load(f)
                if 0 <= idx < len(bookmarks):
                    bookmarks.pop(idx)
                    with open(bm_file, "w") as f:
                        json.dump(bookmarks, f)
                    return make_json({"ok": True})
                return make_json({"ok": False, "error": "Invalid index"})
            except Exception as e:
                return make_json({"ok": False, "error": str(e)})
        return make_json({"error": "Unknown action"})

    return make_response(404, b"Unknown API")


# ── Office Document Parsers ───────────────────────────────────────────

def parse_excel(file_path, size):
    """Parse .xlsx/.xls -> JSON with sheets and table data."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= 500:
                    break
                rows.append([str(cell) if cell is not None else "" for cell in row])
                row_count += 1
            sheets.append({
                "name": sheet_name,
                "rows": rows,
                "total_rows": ws.max_row or 0,
                "total_cols": ws.max_column or 0,
            })
        wb.close()
        return make_json({
            "path": file_path, "size": size,
            "type": "excel",
            "sheets": sheets,
        })
    except Exception as e:
        return make_json({
            "path": file_path, "size": size, "type": "excel",
            "error": f"Excel parse failed: {e}",
        })


def parse_hwp(file_path, size):
    """Parse .hwp -> extracted text."""
    try:
        texts = []

        # Try hwp5txt command first (most reliable)
        try:
            result = subprocess.run(
                ["hwp5txt", file_path],
                capture_output=True, text=True, timeout=10
            )
            if result.stdout.strip():
                texts = [result.stdout]
        except Exception:
            pass

        # Fallback: try python hwp5 library
        if not texts:
            try:
                import hwp5
                from hwp5.xmlmodel import Hwp5File
                hwp5file = Hwp5File(file_path)
                try:
                    for paragraph in hwp5file.bodytext.section(0).paragraphs():
                        texts.append(paragraph.text)
                except Exception:
                    pass
            except Exception:
                pass

        # Last fallback: raw OLE extraction
        if not texts:
            try:
                from hwp5.storage.ole import OleStorage
                olefile = OleStorage(file_path)
                for section_idx in range(20):
                    try:
                        section = olefile.open(f"BodyText/Section{section_idx}")
                        data = section.read()
                        text = data.decode("utf-16-le", errors="replace")
                        # Remove surrogates and non-printable chars
                        clean = "".join(
                            c if (c.isprintable() or c in "\n\r\t") and (ord(c) < 0xD800 or ord(c) > 0xDFFF)
                            else " " for c in text
                        )
                        if clean.strip():
                            texts.append(clean.strip())
                    except Exception:
                        break
            except Exception:
                pass

        content = "\n\n".join(texts) if texts else "[HWP text extraction failed]"

        return make_json({
            "path": file_path, "size": size,
            "type": "hwp",
            "content": content[:200000],
        })
    except Exception as e:
        return make_json({
            "path": file_path, "size": size, "type": "hwp",
            "error": f"HWP parse failed: {e}",
        })


def parse_docx(file_path, size):
    """Parse .docx -> structured content (paragraphs + tables)."""
    try:
        from docx import Document
        doc = Document(file_path)

        elements = []

        for para in doc.paragraphs[:1000]:
            if para.text.strip():
                style = para.style.name if para.style else ""
                elements.append({
                    "type": "paragraph",
                    "text": para.text,
                    "style": style,
                    "bold": any(run.bold for run in para.runs if run.bold),
                })

        for table in doc.tables[:20]:
            rows = []
            for row in table.rows[:200]:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            if rows:
                elements.append({
                    "type": "table",
                    "rows": rows,
                })

        return make_json({
            "path": file_path, "size": size,
            "type": "docx",
            "elements": elements,
        })
    except Exception as e:
        return make_json({
            "path": file_path, "size": size, "type": "docx",
            "error": f"DOCX parse failed: {e}",
        })


# ── Terminal WebSocket ────────────────────────────────────────────────

MAX_SESSIONS = 10
active_sessions = set()


async def terminal_session(websocket):
    # Connection limit
    if len(active_sessions) >= MAX_SESSIONS:
        await websocket.send(json.dumps({
            "type": "output",
            "data": "\r\n\x1b[31m[Max sessions reached -- close another session first]\x1b[0m\r\n"
        }))
        await websocket.close()
        return

    session_id = id(websocket)
    active_sessions.add(session_id)

    # Get session info from the connection (set during process_request)
    session_info = getattr(websocket, "session_info", {"username": "user", "home": FALLBACK_ROOT or str(Path.home())})
    user_home = session_info.get("home", str(Path.home()))

    env = os.environ.copy()
    env["TERM"] = "xterm-256color"
    env["LANG"] = "ko_KR.UTF-8"
    env["HOME"] = user_home

    master_fd = None
    proc = None

    try:
        master_fd, slave_fd = pty.openpty()

        shell = os.environ.get("SHELL", "/bin/bash")
        proc = subprocess.Popen(
            [shell, "--login"],
            stdin=slave_fd, stdout=slave_fd, stderr=slave_fd,
            env=env, preexec_fn=os.setsid, close_fds=True,
            cwd=user_home if os.path.isdir(user_home) else "/tmp",
        )
        os.close(slave_fd)
        slave_fd = -1  # mark as closed

        # Flag to signal shutdown
        shutdown_event = asyncio.Event()

        async def read_pty():
            """Read from PTY using os.read in a thread with select() for timeout."""
            loop = asyncio.get_event_loop()
            try:
                while not shutdown_event.is_set():
                    # Use select with timeout in a thread so we don't block the event loop
                    try:
                        readable = await loop.run_in_executor(
                            None,
                            _select_read, master_fd, 0.5
                        )
                    except (OSError, ValueError):
                        # fd closed or invalid
                        break

                    if readable:
                        try:
                            data = os.read(master_fd, 16384)
                            if not data:
                                break
                            await websocket.send(json.dumps({
                                "type": "output",
                                "data": data.decode("utf-8", errors="replace")
                            }))
                        except OSError:
                            break
                        except websockets.exceptions.ConnectionClosed:
                            break

                    # Check if process exited
                    if proc.poll() is not None:
                        # Read any remaining output
                        try:
                            while True:
                                r = _select_read(master_fd, 0.1)
                                if not r:
                                    break
                                data = os.read(master_fd, 16384)
                                if not data:
                                    break
                                await websocket.send(json.dumps({
                                    "type": "output",
                                    "data": data.decode("utf-8", errors="replace")
                                }))
                        except (OSError, websockets.exceptions.ConnectionClosed):
                            pass
                        break
            except asyncio.CancelledError:
                pass

        async def write_pty():
            """Read from WebSocket and write to PTY."""
            try:
                async for message in websocket:
                    try:
                        msg = json.loads(message)
                    except json.JSONDecodeError:
                        continue
                    if msg.get("type") == "input":
                        try:
                            os.write(master_fd, msg["data"].encode("utf-8"))
                        except OSError:
                            break
                    elif msg.get("type") == "resize":
                        try:
                            rows = msg.get("rows", 24)
                            cols = msg.get("cols", 80)
                            winsize = struct.pack("HHHH", rows, cols, 0, 0)
                            fcntl.ioctl(master_fd, termios.TIOCSWINSZ, winsize)
                        except (OSError, ValueError):
                            pass
            except websockets.exceptions.ConnectionClosed:
                pass
            except asyncio.CancelledError:
                pass

        read_task = asyncio.create_task(read_pty())
        write_task = asyncio.create_task(write_pty())

        # Wait for either task to finish
        done, pending = await asyncio.wait(
            [read_task, write_task],
            return_when=asyncio.FIRST_COMPLETED,
        )

        # Signal shutdown and cancel remaining tasks
        shutdown_event.set()
        for task in pending:
            task.cancel()
            try:
                await task
            except (asyncio.CancelledError, Exception):
                pass

    finally:
        # Kill process group (shell + all children)
        if proc and proc.poll() is None:
            try:
                pgid = os.getpgid(proc.pid)
                os.killpg(pgid, signal.SIGTERM)
            except (ProcessLookupError, OSError):
                pass

            try:
                proc.wait(timeout=2)
            except subprocess.TimeoutExpired:
                try:
                    pgid = os.getpgid(proc.pid)
                    os.killpg(pgid, signal.SIGKILL)
                except (ProcessLookupError, OSError):
                    pass
                try:
                    proc.wait(timeout=1)
                except subprocess.TimeoutExpired:
                    pass

        # Close master fd
        if master_fd is not None:
            try:
                os.close(master_fd)
            except OSError:
                pass

        active_sessions.discard(session_id)


def _select_read(fd, timeout):
    """Blocking select wrapper for use in executor. Returns True if readable."""
    try:
        r, _, _ = select.select([fd], [], [], timeout)
        return bool(r)
    except (OSError, ValueError):
        return False


async def main():
    precompute_tokens()

    auth_desc = {
        "system": "Linux system auth (su)",
        "static": f"static accounts: {', '.join(ACCOUNTS.keys())}",
        "single": f"single password: {FALLBACK_PASSWORD}",
    }[AUTH_MODE]

    print(f"""
+----------------------------------------------------------+
|  Servit — Mobile Server Dashboard                        |
|                                                          |
|  URL:  http://localhost:{PORT:<5}                          |
|  Auth: {auth_desc[:48]:<48} |
|                                                          |
|  Ctrl+C to stop                                          |
+----------------------------------------------------------+
""")

    async with serve(
        terminal_session, HOST, PORT,
        process_request=process_request,
        ping_interval=20,
        ping_timeout=20,
        close_timeout=5,
        max_size=2**21,  # 2MB max message
    ):
        await asyncio.Future()


if __name__ == "__main__":
    asyncio.run(main())
